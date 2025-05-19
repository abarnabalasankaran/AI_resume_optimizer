from flask import Flask, request, jsonify, render_template, send_file
import google.generativeai as genai
from flask_cors import CORS
import io
import magic
import json
import mammoth
import PyPDF2
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import re
import os
app = Flask(__name__)
CORS(app)  


genai.configure(api_key="AIzaSyBLMzqIknsqD0s9WW3OHoiCpbiEdul1vmE")
def extract_text_from_file(file):
    """
    Extract text from uploaded file (PDF or DOCX)
    """

    filename = file.filename.lower()
    
    try:

        mime = magic.Magic(mime=True)
        file_mime = mime.from_buffer(file.read(2048))
        file.seek(0)
        
        if 'pdf' in filename or file_mime == 'application/pdf':

            try:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
                return text
            except Exception as e:
                raise ValueError(f"Error processing PDF: {str(e)}")
                
        elif 'docx' in filename or file_mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':

            try:

                result = mammoth.convert_to_text(file)
                return result.value
            except AttributeError:

                result = mammoth.extract_raw_text(file)
                return result.value
            except Exception as e:
                raise ValueError(f"Error processing DOCX: {str(e)}")
                
        else:

            raise ValueError("Unsupported file type. Please upload PDF or DOCX.")
            
    except Exception as e:

        if filename.endswith('.pdf'):
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
                return text
            except Exception as pdf_err:
                raise ValueError(f"Error processing PDF: {str(pdf_err)}")
                
        elif filename.endswith('.docx'):
            try:

                try:
                    result = mammoth.convert_to_text(file)
                    return result.value
                except AttributeError:

                    result = mammoth.convert_to_html(file)

                    text = re.sub(r'<[^>]+>', ' ', result.value)
                    return re.sub(r'\s+', ' ', text).strip()
            except Exception as docx_err:
                raise ValueError(f"Error processing DOCX: {str(docx_err)}")
                
        else:
            raise ValueError("Unsupported file type. Please upload PDF or DOCX.")

def convert_to_pdf(text):
    """
    Simple PDF conversion for backward compatibility
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                             rightMargin=72, leftMargin=72, 
                             topMargin=36, bottomMargin=18)
    
    styles = getSampleStyleSheet()

    small_style = ParagraphStyle(
        'Small',
        parent=styles['Normal'],
        fontSize=9,
        leading=11
    )
    
    story = []
    

    cleaned_text = text.replace('*', '')

    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    

    paragraphs = cleaned_text.split('\n')
    for para in paragraphs:
        if para.strip():
            p = Paragraph(para, small_style)
            story.append(p)

            story.append(Spacer(1, 3))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_to_pdf_classic(text, template_style):
    """
    Convert text to a downloadable PDF with classic template style
    """
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                             rightMargin=54, leftMargin=54, 
                             topMargin=36, bottomMargin=18)
    

    styles = getSampleStyleSheet()
    
    if template_style == "classic":
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=6,
            textColor=colors.darkblue
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=6,
            spaceAfter=3,
            textColor=colors.darkblue
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,
            leading=11
        )
    
    elif template_style == "modern":
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=6,
            textColor=colors.darkblue,
            alignment=1
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=6,
            spaceAfter=3,
            textColor=colors.teal
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,
            leading=11,
            textColor=colors.black
        )
    
    elif template_style == "minimal":
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=13,
            spaceAfter=6,
            textColor=colors.black
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=3,
            textColor=colors.gray
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,
            leading=11
        )
    
    else:
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=6,
            textColor=colors.black
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=6,
            spaceAfter=3,
            textColor=colors.darkblue
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=9,
            leading=11
        )
    
    story = []
    

    cleaned_text = text.replace('*', '')

    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    

    sections = cleaned_text.split('\n\n')
    

    if sections:
        story.append(Paragraph(sections[0], title_style))
        story.append(Spacer(1, 6))
    

    current_heading = None
    current_content = []
    
    for section in sections[1:]:

        lines = section.split('\n')
        if len(lines) > 0 and (lines[0].isupper() or lines[0].endswith(':') or len(lines[0]) < 30):

            if current_heading:
                story.append(Paragraph(current_heading, heading_style))
                for content in current_content:
                    if content.strip():
                        story.append(Paragraph(content, normal_style))
                story.append(Spacer(1, 3))
            
            current_heading = lines[0]
            current_content = lines[1:] if len(lines) > 1 else []
        else:

            current_content.extend(lines)
    

    if current_heading:
        story.append(Paragraph(current_heading, heading_style))
        for content in current_content:
            if content.strip():
                story.append(Paragraph(content, normal_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def convert_to_docx(text):
    """
    Simple DOCX conversion for backward compatibility
    """
    document = docx.Document()
    

    cleaned_text = text.replace('*', '')

    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    

    style = document.styles['Normal']
    style.font.size = Pt(9)
    

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    

    paragraphs = cleaned_text.split('\n')
    for para in paragraphs:
        if para.strip():
            document.add_paragraph(para)
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def convert_to_docx_template(text, template_style):
    """
    Convert text to a downloadable DOCX with selected template
    """
    document = docx.Document()
    

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    

    if template_style == "classic":

        document.styles['Normal'].font.name = 'Garamond'
        document.styles['Normal'].font.size = Pt(9)
        

        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Garamond'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        

        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Garamond'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        
    elif template_style == "modern":

        document.styles['Normal'].font.name = 'Calibri'
        document.styles['Normal'].font.size = Pt(9)
        

        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        

        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        
    elif template_style == "minimal":

        document.styles['Normal'].font.name = 'Arial'
        document.styles['Normal'].font.size = Pt(9)
        

        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(13)
        title_style.font.bold = True
        

        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(11)
        heading_style.font.bold = True
        
    else:

        document.styles['Normal'].font.name = 'Times New Roman'
        document.styles['Normal'].font.size = Pt(9)
        

        title_style = document.styles['Title'] if 'Title' in document.styles else document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        

        heading_style = document.styles['Heading'] if 'Heading' in document.styles else document.styles.add_style('Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
    

    cleaned_text = text.replace('*', '')

    import re
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    

    sections = cleaned_text.split('\n\n')
    

    if sections:
        title_para = document.add_paragraph(style='Title')
        title_para.add_run(sections[0])
        if template_style == "modern":
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_para.paragraph_format.space_after = Pt(6)
    

    current_heading = None
    current_content = []
    
    for section in sections[1:]:

        lines = section.split('\n')
        if len(lines) > 0 and (lines[0].isupper() or lines[0].endswith(':') or len(lines[0]) < 30):

            if current_heading:
                head_para = document.add_paragraph(style='Heading')
                head_para.add_run(current_heading)

                head_para.paragraph_format.space_before = Pt(6)
                head_para.paragraph_format.space_after = Pt(3)
                
                for content in current_content:
                    if content.strip():
                        para = document.add_paragraph(content, style='Normal')

                        para.paragraph_format.space_after = Pt(0)
            
            current_heading = lines[0]
            current_content = lines[1:] if len(lines) > 1 else []
        else:

            current_content.extend(lines)
    

    if current_heading:
        head_para = document.add_paragraph(style='Heading')
        head_para.add_run(current_heading)

        head_para.paragraph_format.space_before = Pt(6)
        head_para.paragraph_format.space_after = Pt(3)
        
        for content in current_content:
            if content.strip():
                para = document.add_paragraph(content, style='Normal')

                para.paragraph_format.space_after = Pt(0)
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def extract_skills(job_description):
    """
    Extract skills from job description and categorize them
    """
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"""
    Analyze the following job description and extract skills in these categories:
    1. Technical Skills: programming languages, tools, platforms, etc.
    2. Soft Skills: communication, teamwork, leadership, etc.
    3. Domain Knowledge: industry-specific knowledge, regulations, etc.
    
    Format your response as JSON with these three categories as keys and arrays of skills as values.
    
    Job Description:
    {job_description}
    """
    
    response = model.generate_content(prompt)
    skills_text = response.text.strip()
    

    import json
    try:

        json_match = re.search(r'```json(.*?)```', skills_text, re.DOTALL)
        if json_match:
            skills_json = json.loads(json_match.group(1).strip())
        else:

            skills_json = json.loads(skills_text)
            

        all_skills = []
        for category in skills_json:
            all_skills.extend(skills_json[category])
            
        return skills_json, all_skills
        
    except (json.JSONDecodeError, AttributeError):

        skills = skills_text.split(",")
        skills = [skill.strip() for skill in skills]

        skills_json = {
            "Technical Skills": skills,
            "Soft Skills": [],
            "Domain Knowledge": []
        }
        return skills_json, skills

def analyze_resume(resume_text, skills, skills_by_category=None):
    """
    Enhanced resume analysis with matching by category and detailed feedback
    """

    if skills_by_category is None:
        skills_by_category = {
            "Technical Skills": skills,
            "Soft Skills": [],
            "Domain Knowledge": []
        }
    

    matched_skills = [skill for skill in skills if skill.lower() in resume_text.lower()]
    match_percentage = (len(matched_skills) / len(skills)) * 100 if skills else 0
    

    category_analysis = {}
    for category, category_skills in skills_by_category.items():
        if not category_skills:
            category_analysis[category] = {
                "matched": [],
                "unmatched": [],
                "match_percentage": 0
            }
            continue
            
        matched = [skill for skill in category_skills if skill.lower() in resume_text.lower()]
        unmatched = [skill for skill in category_skills if skill.lower() not in resume_text.lower()]
        match_pct = (len(matched) / len(category_skills)) * 100 if category_skills else 0
        
        category_analysis[category] = {
            "matched": matched,
            "unmatched": unmatched,
            "match_percentage": match_pct
        }
    

    model = genai.GenerativeModel("gemini-1.5-flash")
    

    skills_text = ""
    for category, category_skills in skills_by_category.items():
        skills_text += f"{category}:\n"
        for skill in category_skills:
            match_status = "✓" if skill.lower() in resume_text.lower() else "✗"
            skills_text += f"- {match_status} {skill}\n"
    
    prompt = f"""
    Analyze this resume against the required job skills:
    
    RESUME:
    {resume_text}
    
    REQUIRED SKILLS:
    {skills_text}
    
    Provide a detailed analysis with:
    1. Specific recommendations to improve match rate
    2. Priority skills to add or emphasize
    3. Sections that need improvement
    
    Format your response as JSON with these keys: 
    "recommendations", "priority_skills", "sections_to_improve"
    """
    
    try:
        response = model.generate_content(prompt)
        analysis_text = response.text.strip()
        

        json_match = re.search(r'```json(.*?)```', analysis_text, re.DOTALL)
        if json_match:
            detailed_analysis = json.loads(json_match.group(1).strip())
        else:

            detailed_analysis = json.loads(analysis_text)
    except:

        detailed_analysis = {
            "recommendations": ["Ensure your resume highlights relevant skills explicitly"],
            "priority_skills": [skill for skill in skills if skill not in matched_skills][:3],
            "sections_to_improve": ["Skills section", "Work experience"]
        }
    

    if match_percentage < 40:
        emotion = "Needs improvement"
    elif match_percentage < 70:
        emotion = "Good potential"
    else:
        emotion = "Excellent match!"
    
    return {
        "matched_skills": matched_skills,
        "unmatched_skills": [skill for skill in skills if skill not in matched_skills],
        "match_percentage": match_percentage,
        "emotion": emotion,
        "category_analysis": category_analysis,
        "detailed_analysis": detailed_analysis
    }

def tailor_resume(resume_text, job_description):
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"Given this resume:\n{resume_text}\nand this job description:\n{job_description}\n" \
             "Rewrite the resume to highlight relevant skills and experience, focusing on:" \
             "1. Matching job requirements\n" \
             "2. Emphasizing transferable skills\n" \
             "3. Using industry-specific keywords\n" \
             "4. Maintaining professional tone\n" \
             "5. Structure the resume with clear sections for: Contact Information, Professional Summary, Work Experience, Skills, Education.\n" \
             "6. Format each job entry with company, title, dates, and bullet points for achievements.\n" \
             "7. IMPORTANT: Keep the resume concise to fit on ONE page. Do not use asterisk (*) symbols for bullet points - use hyphens (-) instead.\n" \
             "8. Avoid leaving large gaps between sections and keep descriptions brief but impactful.\n" \
             "9. Use periods at the end of achievement statements only if they form complete sentences."

    response = model.generate_content(prompt)
    return response.text.strip()

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/extract_skills', methods=['POST'])
def get_skills():
    data = request.get_json()
    job_description = data.get('job_description', '')
    skills_by_category, all_skills = extract_skills(job_description)
    return jsonify({
        "skills": all_skills,
        "skills_by_category": skills_by_category
    })

@app.route('/upload_resume', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    try:
        resume_text = extract_text_from_file(file)
        return jsonify({"resume_text": resume_text})
    except Exception as e:
        return jsonify({"error": str(e)}), 400

@app.route('/analyze_resume', methods=['POST'])
def analyze():
    data = request.get_json()
    resume_text = data.get('resume_text', '')
    skills = data.get('skills', [])
    skills_by_category = data.get('skills_by_category', None)
    
    analysis_result = analyze_resume(resume_text, skills, skills_by_category)
    return jsonify(analysis_result)

@app.route('/tailor_resume', methods=['POST'])
def tailor():

    resume_text = request.form.get('resume_text', '')
    job_description = request.form.get('job_description', '')
    output_format = request.form.get('output_format', 'pdf')
    template_style = request.form.get('template_style', 'professional')
    
    tailored_resume = tailor_resume(resume_text, job_description)
    
    if output_format == 'docx':
        docx_file = convert_to_docx_template(tailored_resume, template_style)
        return send_file(docx_file, 
                         download_name='tailored_resume.docx', 
                         as_attachment=True, 
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        pdf_file = convert_to_pdf_classic(tailored_resume, template_style)
        return send_file(pdf_file, 
                         download_name='tailored_resume.pdf', 
                         as_attachment=True, 
                         mimetype='application/pdf')

@app.route('/preview_resume', methods=['POST'])
def preview_resume():
    data = request.get_json()
    resume_text = data.get('resume_text', '')
    job_description = data.get('job_description', '')
    
    tailored_resume = tailor_resume(resume_text, job_description)
    return jsonify({"tailored_resume": tailored_resume})

if __name__ == '__main__':
    app.run(debug=True)